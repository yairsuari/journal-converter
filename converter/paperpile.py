"""
Extract and re-inject Paperpile citation fields in .docx files.

Paperpile stores citation data as Word ADDIN fields with a base64+zlib-compressed
JSON payload containing BibTeX citekeys.  This module provides two services:

1. extract_citation_map()  — builds {rendered_text → citekeys} for cite_revive
   (DOCX → LaTeX path: replace plain text with \\citep{key})

2. inject_paperpile_fields() — restores live Paperpile field codes in an output
   .docx, so the converted document still has working citations in Word
   (DOCX → DOCX path: re-inject ADDIN fields that Pandoc discarded)
"""
from __future__ import annotations
import base64
import json
import re
import zlib
from copy import deepcopy
from pathlib import Path

_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W = f'{{{_NS}}}'
_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


# ── XML helpers ───────────────────────────────────────────────────────────────

def _decode_data(b64: str) -> list[dict]:
    padded = b64 + '=' * (4 - len(b64) % 4)
    return json.loads(zlib.decompress(base64.b64decode(padded)))


def _iter_fields(para_elem):
    """Yield (instr_text, rendered_text) for each Word field in a paragraph element."""
    in_f = past_sep = False
    instr: list[str] = []
    render: list[str] = []
    for child in para_elem.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'fldChar':
            ft = child.get(f'{_W}fldCharType')
            if ft == 'begin':
                in_f, past_sep, instr, render = True, False, [], []
            elif ft == 'separate':
                past_sep = True
            elif ft == 'end' and in_f:
                yield ''.join(instr), ''.join(render)
                in_f = False
        elif tag == 'instrText' and in_f and not past_sep:
            instr.append(child.text or '')
        elif tag == 't' and in_f and past_sep:
            render.append(child.text or '')


def _normalize(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def _make_field_runs(instr_text: str, rendered: str, rpr=None) -> list:
    """Return the 5-run sequence that encodes a Word ADDIN field."""
    from lxml import etree  # noqa: PLC0415

    def run(*children):
        r = etree.Element(f'{_W}r')
        if rpr is not None:
            r.append(deepcopy(rpr))
        for c in children:
            r.append(c)
        return r

    fc_begin = etree.Element(f'{_W}fldChar')
    fc_begin.set(f'{_W}fldCharType', 'begin')

    instr_elem = etree.Element(f'{_W}instrText')
    instr_elem.set(_XML_SPACE, 'preserve')
    instr_elem.text = instr_text

    fc_sep = etree.Element(f'{_W}fldChar')
    fc_sep.set(f'{_W}fldCharType', 'separate')

    t_elem = etree.Element(f'{_W}t')
    t_elem.set(_XML_SPACE, 'preserve')
    t_elem.text = rendered

    fc_end = etree.Element(f'{_W}fldChar')
    fc_end.set(f'{_W}fldCharType', 'end')

    return [run(fc_begin), run(instr_elem), run(fc_sep), run(t_elem), run(fc_end)]


def _text_run(template_run, text: str):
    """Create a plain-text run, copying rPr from template_run."""
    from lxml import etree  # noqa: PLC0415
    r = etree.Element(f'{_W}r')
    rpr = template_run.find(f'{_W}rPr')
    if rpr is not None:
        r.append(deepcopy(rpr))
    t = etree.Element(f'{_W}t')
    t.text = text
    if text != text.strip():
        t.set(_XML_SPACE, 'preserve')
    r.append(t)
    return r


def _inject_in_para(p_elem, field_map: dict[str, tuple[str, str]]) -> int:
    """
    Inject Paperpile fields into a single paragraph element.
    Restarts the scan after each injection so that a run containing two
    citations (e.g. after a prior split) is handled correctly.
    Returns the number of fields injected.
    """
    count = 0
    while True:
        injected = False
        for child in list(p_elem):
            if child.tag != f'{_W}r':
                continue
            if (child.find(f'{_W}fldChar') is not None or
                    child.find(f'{_W}instrText') is not None):
                continue
            t_elem = child.find(f'{_W}t')
            if t_elem is None or not t_elem.text:
                continue

            run_text = t_elem.text
            for orig_rendered, instr_text in field_map.values():
                pos = run_text.find(orig_rendered)
                if pos == -1:
                    continue

                end_pos = pos + len(orig_rendered)
                rpr = child.find(f'{_W}rPr')

                replacements: list = []
                if run_text[:pos]:
                    replacements.append(_text_run(child, run_text[:pos]))
                replacements.extend(_make_field_runs(instr_text, orig_rendered, rpr))
                if run_text[end_pos:]:
                    replacements.append(_text_run(child, run_text[end_pos:]))

                idx = list(p_elem).index(child)
                p_elem.remove(child)
                for j, r in enumerate(replacements):
                    p_elem.insert(idx + j, r)

                count += 1
                injected = True
                break
            if injected:
                break  # restart paragraph scan
        if not injected:
            break
    return count


# ── Public API ────────────────────────────────────────────────────────────────

def extract_citation_map(docx_path: Path) -> dict[str, list[str]]:
    """
    Return {normalized_rendered_text: [citekey, ...]} for every Paperpile field.

    Keys include outer parentheses, e.g. '(Smith, 2020)' -> ['Smith2020-ab'].
    Multi-citation fields produce a single entry:
      '(Smith, 2020; Jones, 2021)' -> ['Smith2020-ab', 'Jones2021-xy']
    Returns {} when the source is not a .docx or has no Paperpile fields.
    Used by cite_revive for the DOCX → LaTeX path.
    """
    if docx_path.suffix.lower() != '.docx':
        return {}
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return {}
    try:
        doc = Document(docx_path)
    except Exception:
        return {}

    result: dict[str, list[str]] = {}
    for para in doc.paragraphs:
        for instr, rendered in _iter_fields(para._p):
            if 'ADDIN paperpile_citation' not in instr:
                continue
            m = re.search(r'<data>(.*?)</data>', instr, re.DOTALL)
            if not m:
                continue
            try:
                citations = _decode_data(m.group(1))
            except Exception:
                continue
            keys = [c.get('citekey') for c in citations if c.get('citekey')]
            if keys and rendered.strip():
                result[_normalize(rendered)] = keys
    return result


def extract_field_map(docx_path: Path) -> dict[str, tuple[str, str]]:
    """
    Return {normalized_rendered: (original_rendered, instr_text)} for every
    Paperpile field in docx_path.
    Used by inject_paperpile_fields for the DOCX → DOCX path.
    Returns {} when not applicable.
    """
    if docx_path.suffix.lower() != '.docx':
        return {}
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return {}
    try:
        doc = Document(docx_path)
    except Exception:
        return {}

    result: dict[str, tuple[str, str]] = {}
    for para in doc.paragraphs:
        for instr, rendered in _iter_fields(para._p):
            if 'ADDIN paperpile_citation' not in instr:
                continue
            if rendered.strip():
                norm = _normalize(rendered)
                if norm not in result:
                    result[norm] = (rendered, instr)
    return result


def inject_paperpile_fields(docx_path: Path,
                             field_map: dict[str, tuple[str, str]]) -> int:
    """
    Post-process a .docx file to replace plain-text citation strings with live
    Paperpile ADDIN field codes.

    field_map: output of extract_field_map() from the *source* document.
    Returns the number of fields injected.  Modifies docx_path in place.
    """
    if not field_map:
        return 0
    try:
        from docx import Document  # noqa: PLC0415
        from docx.oxml.ns import qn  # noqa: PLC0415
    except ImportError:
        return 0
    try:
        doc = Document(docx_path)
    except Exception:
        return 0

    count = 0
    for p_elem in doc.element.body.iter(qn('w:p')):
        count += _inject_in_para(p_elem, field_map)

    if count > 0:
        doc.save(docx_path)
    return count
